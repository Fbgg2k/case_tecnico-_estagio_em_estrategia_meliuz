import os
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import markdown
from bs4 import BeautifulSoup

def convert_md_to_docx(input_md_path, output_docx_path):
    # Read the markdown file
    with open(input_md_path, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert markdown to HTML
    html = markdown.markdown(md_content)
    
    # Create a new Word document
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Parse HTML with BeautifulSoup
    soup = BeautifulSoup(html, 'html.parser')
    
    # Process each HTML element
    for element in soup.find_all(True):
        if element.name == 'h1':
            paragraph = doc.add_heading(level=1)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run(element.get_text())
            run.bold = True
            run.font.size = Pt(16)
        elif element.name == 'h2':
            paragraph = doc.add_heading(level=2)
            run = paragraph.add_run(element.get_text())
            run.bold = True
            run.font.size = Pt(14)
        elif element.name == 'h3':
            paragraph = doc.add_heading(level=3)
            run = paragraph.add_run(element.get_text())
            run.bold = True
            run.italic = True
            run.font.size = Pt(12)
        elif element.name == 'p':
            doc.add_paragraph(element.get_text())
        elif element.name == 'ul':
            for li in element.find_all('li'):
                paragraph = doc.add_paragraph(style='List Bullet')
                paragraph.add_run(li.get_text())
        elif element.name == 'ol':
            for i, li in enumerate(element.find_all('li'), 1):
                paragraph = doc.add_paragraph(style='List Number')
                paragraph.add_run(f"{i}. {li.get_text()}")
        elif element.name == 'strong':
            run = doc.add_paragraph().add_run(element.get_text())
            run.bold = True
        elif element.name == 'em':
            run = doc.add_paragraph().add_run(element.get_text())
            run.italic = True
        elif element.name == 'code':
            run = doc.add_paragraph().add_run(element.get_text())
            run.font.name = 'Courier New'
    
    # Save the document
    doc.save(output_docx_path)

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python convert_to_docx.py <input_md_file> <output_docx_file>")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2]
    
    if not os.path.exists(input_file):
        print(f"Error: Input file '{input_file}' not found.")
        sys.exit(1)
    
    try:
        convert_md_to_docx(input_file, output_file)
        print(f"Successfully converted {input_file} to {output_file}")
    except Exception as e:
        print(f"Error during conversion: {str(e)}")
        sys.exit(1)
